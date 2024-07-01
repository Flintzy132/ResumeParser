from flask import Flask, request, redirect, url_for, render_template
import os
from PyPDF2 import PdfReader
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document

app = Flask(__name__)

load_dotenv()

client = OpenAI(
    api_key=os.environ.get('OPENAI_API_KEY'),
)


@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'resumes[]' not in request.files or 'job_description' not in request.files:
        return redirect(url_for('index'))

    resume_files = request.files.getlist('resumes[]')
    jd_file = request.files['job_description']

    if any(f.filename == '' for f in resume_files) or jd_file.filename == '':
        return redirect(url_for('index'))

    resumes_content = ''
    for resume in resume_files:
        reader = PdfReader(resume)
        for page in reader.pages:
            resumes_content += page.extract_text()

    jd_content = ''
    reader1 = PdfReader(jd_file)
    for page in reader1.pages:
        jd_content += page.extract_text()

    combined_prompt = (
        f"Resumes content:\n{resumes_content}\n\n"
        f"JD content:\n{jd_content}\n\n"
        "Question: Please analyze the all candidates' resumes in relation to the job description. Identify skills that align between the resumes and the jd documents and provide insights on whether these candidates should be considered for the job based on these matches. Mention which candidate you are writing for by mentioning their name and analyse all of them seperately, in the final statement also mention which candidate are relevant and which is nots"
    )
    response1 = ask_openai(combined_prompt)
    doc = Document()
    doc.add_heading('Result', level=1)
    doc.add_paragraph(response1)
    doc.save('ResumeInsights.docx')
    return render_template('index.html', response=response1)


def ask_openai(prompt):
    try:
        response = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": prompt,
                }
            ],
            model="gpt-3.5-turbo",
        )
        return response.choices[0].message.content
    except Exception as e:
        return str(e)


if __name__ == '__main__':
    app.run(debug=True, port=8080)
